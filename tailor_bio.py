#!/usr/bin/env python3
"""Synaptiq Bio Tailoring CLI — researches a prospect company and generates
a tailored consultant bio, company profile, tailoring rationale, and
recommended discovery questions.

Outputs:
  1. Tailored bio (.docx) with headshot and logo from the original
  2. Tailoring specifics (.md) — what was changed and why
  3. Company research profile (.md)
  4. Recommended discovery questions (.md)
"""

import argparse
import json
import logging
import os
import re
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree as lxml_etree
from dotenv import load_dotenv

load_dotenv(override=True)

SCRIPT_DIR = Path(__file__).resolve().parent
LOG = logging.getLogger("tailor_bio")

DEFAULT_MODEL = "claude-sonnet-4-20250514"
RESEARCH_MODEL = "claude-sonnet-4-20250514"


# ---------------------------------------------------------------------------
# Bio extraction from .docx
# ---------------------------------------------------------------------------

def extract_bio_text(docx_path: Path) -> str:
    """Extract all text from a .docx file using zipfile + XML parsing."""
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(str(docx_path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        # Fallback: raw XML extraction
        z = zipfile.ZipFile(str(docx_path))
        tree = ET.parse(z.open("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for p in tree.findall(".//w:p", ns):
            texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n\n".join(paragraphs)


def extract_bio_images(docx_path: Path, tmp_dir: Path) -> dict[str, Path]:
    """Extract images from a .docx and return a dict of name -> path."""
    images = {}
    z = zipfile.ZipFile(str(docx_path))
    for name in z.namelist():
        if "media/" in name:
            extracted = Path(z.extract(name, str(tmp_dir)))
            images[Path(name).name] = extracted
    return images


# ---------------------------------------------------------------------------
# Claude API helpers
# ---------------------------------------------------------------------------

def call_claude(
    client: anthropic.Anthropic,
    system: str,
    user_message: str,
    model: str,
    max_tokens: int = 8192,
    tools: list | None = None,
) -> str:
    """Call Claude and return the concatenated text response, handling tool use loops."""
    messages = [{"role": "user", "content": user_message}]
    kwargs = dict(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=messages,
    )
    if tools:
        kwargs["tools"] = tools

    while True:
        response = client.messages.create(**kwargs)

        # Collect all text blocks from the response
        text_parts = []
        tool_results = []
        has_tool_use = False

        for block in response.content:
            if block.type == "text":
                text_parts.append(block.text)
            elif block.type == "web_search_tool_result":
                # Server-side web search results are handled automatically
                pass
            elif block.type == "tool_use":
                has_tool_use = True

        # If stop reason is end_turn or no tool use, we're done
        if response.stop_reason == "end_turn" or not has_tool_use:
            return "\n".join(text_parts).strip()

        # Otherwise, continue the conversation with tool results
        messages.append({"role": "assistant", "content": response.content})
        # For server-side tools, the API handles tool execution internally
        # so we just need to check if we got our final text
        if text_parts:
            return "\n".join(text_parts).strip()

        # Safety: if we somehow loop without text, break
        break

    return "\n".join(text_parts).strip() if text_parts else ""


def research_company(
    client: anthropic.Anthropic,
    company_name: str,
    inquiry: str,
    model: str,
) -> str:
    """Use Claude with web search to research the prospect company."""
    LOG.info("Researching %s via web search...", company_name)

    system = (
        "You are a business research analyst. Your job is to compile a thorough "
        "company profile for a prospect company that an AI consulting firm is about "
        "to engage with. The profile should help a consultant understand the company's "
        "culture, structure, technology posture, and strategic priorities so they can "
        "tailor their approach.\n\n"
        "Use web search extensively to gather current, accurate information. "
        "Return a well-structured markdown document."
    )

    user_msg = f"""Research **{company_name}** thoroughly and produce a comprehensive company profile in markdown format.

The context: a contact at {company_name} has submitted the following inquiry to an AI consulting firm:

---
{inquiry}
---

Structure your research profile with these sections:

## Company Overview
- What the company does, size, revenue, headquarters, leadership
- Ownership structure (public/private, family-owned, etc.)

## Culture & Values
- Core values, mission statement, leadership philosophy
- How the company approaches change and innovation
- Employee/talent philosophy

## Workforce & Organizational Structure
- Corporate structure, key divisions/teams
- Any known analytics, data science, or performance insights teams
- Workforce size and growth trajectory

## Technology & AI Posture
- Current technology initiatives (AI, automation, digital transformation)
- How the company has publicly discussed AI adoption
- Any known partnerships or vendor relationships in tech/AI space

## Industry Context
- Industry trends affecting this company
- Competitive landscape relevant to AI adoption
- Regulatory or compliance considerations

## Relevance to the Inquiry
- How the company's culture and values connect to the inquiry themes
- Known challenges or opportunities related to the inquiry
- Key considerations for an AI consultant engaging with this organization

Be thorough and cite specific facts. Focus on information that would help a consultant prepare for a workforce transformation and AI strategy engagement."""

    return call_claude(
        client, system, user_msg, model,
        max_tokens=16384,
        tools=[{"type": "web_search_20250305", "name": "web_search", "max_uses": 10}],
    )


def generate_tailored_bio(
    client: anthropic.Anthropic,
    original_bio: str,
    company_profile_md: str,
    synaptiq_profile: str,
    inquiry: str,
    company_name: str,
    model: str,
) -> dict[str, str]:
    """Generate tailored bio content, tailoring notes, and discovery questions."""
    LOG.info("Generating tailored bio content...")

    system = (
        "You are an expert at crafting executive bios for consulting engagements. "
        "You tailor bios to resonate with specific prospects by emphasizing relevant "
        "experience, aligning language with the prospect's culture and values, and "
        "foregrounding capabilities that directly address the prospect's stated needs. "
        "You never fabricate experience — you reframe and emphasize what's genuinely there.\n\n"
        "Return your response as a JSON object with exactly these keys:\n"
        '  "succinct_bio": a 1-paragraph tailored bio (~100-150 words)\n'
        '  "extended_bio": a multi-paragraph tailored bio (4-5 paragraphs)\n'
        '  "relevant_experience": a JSON array of objects with "title" and "description" keys '
        "(6 bullet points of relevant experience highlights)\n"
        '  "tailoring_notes": a markdown document explaining what you tailored and why\n'
        '  "discovery_questions": a markdown document with 3-5 recommended questions for the initial consult\n\n'
        "Return ONLY valid JSON. No markdown fences."
    )

    user_msg = f"""## Task
Tailor the following bio for an engagement with **{company_name}** based on their inquiry and the research profile below.

## Original Bio
{original_bio}

## Synaptiq Company Profile
{synaptiq_profile}

## Research Profile on {company_name}
{company_profile_md}

## Inquiry from {company_name} Contact
{inquiry}

## Instructions

### For the tailored bios (succinct_bio, extended_bio):
- Reframe Stephen's experience to emphasize what's most relevant to {company_name}'s inquiry
- Align language and themes with {company_name}'s known culture and values
- Emphasize workforce transformation, AI strategy, and organizational design capabilities
- Highlight industry-relevant experience
- Maintain authenticity — only reframe what's genuinely in the original bio
- Connect Synaptiq's philosophy to {company_name}'s values where natural

### For relevant_experience (6 items):
- Each should have a "title" (short capability name) and "description" (1-2 sentences)
- Prioritize experience directly relevant to the inquiry themes
- Include industry-specific experience if applicable

### For tailoring_notes:
- Structure as a markdown document with clear sections
- Explain each major tailoring decision and why it was made
- Reference specific {company_name} cultural/strategic factors that informed choices
- Note what was emphasized, de-emphasized, and any language/framing choices

### For discovery_questions:
- Provide 3-5 questions to ask the contact at the start of a 1-hour consult
- Each question should fill a gap in understanding based on the inquiry
- Questions should demonstrate preparation and research knowledge
- Structure as markdown with the question, why it matters, and what the answer unlocks
- Focus on making the first hour maximally productive"""

    raw = call_claude(client, system, user_msg, model, max_tokens=16384)

    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw[:-3].strip()

    return json.loads(raw)


def generate_prep_doc(
    client: anthropic.Anthropic,
    company_profile_md: str,
    synaptiq_profile: str,
    inquiry: str,
    company_name: str,
    discovery_questions: str,
    tailoring_notes: str,
    model: str,
) -> str:
    """Generate a concise consultant prep document synthesizing all research."""
    LOG.info("Generating consultant prep doc...")

    system = (
        "You are a senior consulting strategist writing internal prep notes for Stephen Sklarew, "
        "CEO of Synaptiq (an AI consulting firm), before a 1-hour introductory consult with a prospect. "
        "Write in a candid, direct, internal tone — not polished or client-facing. "
        "Think of this as the notes a sharp chief of staff would put on Stephen's desk 30 minutes "
        "before the call. Be blunt where useful. Flag landmines. Suggest power moves.\n\n"
        "Return the prep doc as markdown. Target 2-3 pages when printed."
    )

    user_msg = f"""## Context

Stephen has a 1-hour introductory consult with a contact at **{company_name}**. Write his internal prep brief.

## Company Research
{company_profile_md}

## The Inquiry
{inquiry}

## Synaptiq Capabilities
{synaptiq_profile}

## Discovery Questions Already Prepared
{discovery_questions}

## Tailoring Notes (what we emphasized in the bio we sent)
{tailoring_notes}

## Prep Doc Structure

### 1. TL;DR (3-5 bullet executive summary)
- Who they are, what they want, why they're reaching out now
- The one thing Stephen must understand going in

### 2. Company Snapshot
- Distill the research into the key facts Stephen needs (not everything — just what matters for this call)
- Culture/values context that should shape how he shows up
- Internal language or concepts to mirror (e.g., if they call HQ the "Support Center," use that)

### 3. The Inquiry — What's Really Going On
- Restate the inquiry in plain language
- Read between the lines: what's the underlying anxiety or strategic tension?
- What's the contact likely hoping to hear vs. afraid of hearing?

### 4. Key Talking Points & Insights to Bring
- 4-6 specific things Stephen should weave into the conversation
- For each: the insight, why it lands, and how to introduce it naturally
- Include relevant Synaptiq experience/case studies to reference (be specific)
- Flag any insider knowledge from research that shows preparation without being creepy

### 5. Landmines & Things to Avoid
- Topics, framing, or language that could backfire given their culture
- Common consulting missteps for this type of engagement
- Anything from research that's sensitive (e.g., layoffs, controversies, internal politics)

### 6. Recommended Meeting Flow (1 hour)
- Minute-by-minute suggested agenda
- When to listen vs. when to present
- How to open, when to pivot, and how to close with a clear next step

### 7. What Success Looks Like
- **For the contact (this call):** What should they feel/know/believe when they hang up?
- **For the broader engagement:** What outcome are they ultimately trying to achieve?
- **For Stephen/Synaptiq:** What's the ideal next step to propose?

Keep it tight. No filler. This is a working document, not a deliverable."""

    return call_claude(client, system, user_msg, model, max_tokens=8192)


# ---------------------------------------------------------------------------
# Document generation (.docx)
# ---------------------------------------------------------------------------

def build_docx(
    bio_data: dict[str, str],
    images: dict[str, Path],
    output_path: Path,
):
    """Build the tailored bio .docx with headshot and logo."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # --- Logo ---
    logo_path = _find_image(images, ["logo", "image1", ".png"])
    if logo_path:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(3.5))
        logo_para.space_after = Pt(4)

    # --- Divider ---
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = divider.add_run("_" * 70)
    div_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    div_run.font.size = Pt(8)
    divider.space_after = Pt(12)

    # --- Headshot + Name table ---
    headshot_path = _find_image(images, ["headshot", "photo", "image2", ".jpg", ".jpeg"])
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.0)

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders_xml = (
                '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            )
            for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                borders_xml += f'<w:{border} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            borders_xml += "</w:tcBorders>"
            tcPr.append(lxml_etree.fromstring(borders_xml))

    # Headshot
    if headshot_path:
        cell_img = table.cell(0, 0)
        cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = cell_img.paragraphs[0].add_run()
        img_run.add_picture(str(headshot_path), width=Inches(1.5))

    # Name/title
    cell_text = table.cell(0, 1)
    cell_text.vertical_alignment = 1  # CENTER
    cell_text.paragraphs[0].text = ""
    name_run = cell_text.paragraphs[0].add_run("Stephen Sklarew")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    cell_text.paragraphs[0].space_after = Pt(2)

    title_para = cell_text.add_paragraph()
    title_run = title_para.add_run("CEO & Co-Founder, Synaptiq")
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    title_run.italic = True
    title_para.space_after = Pt(4)

    tagline_para = cell_text.add_paragraph()
    tagline_run = tagline_para.add_run(
        'AI Strategy & Workforce Transformation  |  "The Humankind of AI"'
    )
    tagline_run.font.size = Pt(10)
    tagline_run.font.color.rgb = RGBColor(0x66, 0x88, 0xAA)
    tagline_para.space_after = Pt(0)

    # --- Spacer ---
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(6)
    spacer.space_after = Pt(2)

    # --- Overview ---
    _add_heading(doc, "Overview")
    succinct = doc.add_paragraph()
    succinct.add_run(bio_data["succinct_bio"]).font.size = Pt(11)
    succinct.space_after = Pt(12)

    # --- Extended bio ---
    _add_heading(doc, "Background & Expertise")
    for para_text in bio_data["extended_bio"].split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        para = doc.add_paragraph()
        para.add_run(para_text).font.size = Pt(11)
        para.space_after = Pt(8)

    # --- Relevant experience ---
    _add_heading(doc, "Relevant Experience")
    experience = bio_data["relevant_experience"]
    if isinstance(experience, str):
        experience = json.loads(experience)
    for item in experience:
        bullet = doc.add_paragraph()
        bullet.style = "List Bullet"
        title_run = bullet.add_run(f'{item["title"]}: ')
        title_run.bold = True
        title_run.font.size = Pt(10.5)
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        desc_run = bullet.add_run(item["description"])
        desc_run.font.size = Pt(10.5)
        bullet.space_after = Pt(4)

    doc.save(str(output_path))
    LOG.info("Wrote tailored bio: %s", output_path)


def _add_heading(doc, text: str):
    heading = doc.add_paragraph()
    h_run = heading.add_run(text)
    h_run.bold = True
    h_run.font.size = Pt(13)
    h_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    heading.space_after = Pt(6)


def _find_image(images: dict[str, Path], hints: list[str]) -> Path | None:
    """Find an image matching any of the hint strings."""
    for hint in hints:
        for name, path in images.items():
            if hint.lower() in name.lower():
                return path
    # Fallback: return first image if only hints don't match
    return None


# ---------------------------------------------------------------------------
# Markdown output writers
# ---------------------------------------------------------------------------

def write_tailoring_notes(notes: str, company_name: str, output_path: Path):
    header = (
        f"# Bio Tailoring Specifics — {company_name}\n\n"
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n\n"
    )
    output_path.write_text(header + notes, encoding="utf-8")
    LOG.info("Wrote tailoring notes: %s", output_path)


def write_company_profile(profile: str, company_name: str, output_path: Path):
    header = (
        f"# Company Research Profile — {company_name}\n\n"
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n\n"
    )
    output_path.write_text(header + profile, encoding="utf-8")
    LOG.info("Wrote company profile: %s", output_path)


def write_discovery_questions(questions: str, company_name: str, output_path: Path):
    header = (
        f"# Recommended Discovery Questions — {company_name}\n\n"
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n\n"
    )
    output_path.write_text(header + questions, encoding="utf-8")
    LOG.info("Wrote discovery questions: %s", output_path)


def write_prep_doc(prep: str, company_name: str, output_path: Path):
    header = (
        f"# Consultant Prep Brief — {company_name}\n\n"
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n"
        f"*INTERNAL — NOT FOR DISTRIBUTION*\n\n---\n\n"
    )
    output_path.write_text(header + prep, encoding="utf-8")
    LOG.info("Wrote prep doc: %s", output_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def slugify(text: str) -> str:
    """Convert text to a filesystem-safe slug."""
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    return text.strip("-")


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Tailor a Synaptiq consultant bio for a prospect inquiry.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --company "Chick-fil-A" --inquiry "Seeking AI workforce planning guidance..." \\
           --bio "2026 Sklarew Bio - IIA.docx"

  %(prog)s -c "Acme Corp" -q inquiry.txt -b bio.docx -o output/acme
        """,
    )
    p.add_argument(
        "-c", "--company",
        required=True,
        help="Name of the prospect company",
    )
    p.add_argument(
        "-q", "--inquiry",
        required=True,
        help="Inquiry description text, or path to a .txt/.md file containing it",
    )
    p.add_argument(
        "-b", "--bio",
        default=str(SCRIPT_DIR / "2026 Sklarew Bio - IIA.docx"),
        help="Path to the original bio .docx file (default: 2026 Sklarew Bio - IIA.docx)",
    )
    p.add_argument(
        "-p", "--company-profile",
        default=str(SCRIPT_DIR / "synaptiq-company-profile.md"),
        help="Path to the Synaptiq company profile .md (default: synaptiq-company-profile.md)",
    )
    p.add_argument(
        "-o", "--output-dir",
        default=None,
        help="Output directory (default: output/<company-slug>)",
    )
    p.add_argument(
        "-m", "--model",
        default=DEFAULT_MODEL,
        help=f"Claude model for content generation (default: {DEFAULT_MODEL})",
    )
    p.add_argument(
        "--research-model",
        default=RESEARCH_MODEL,
        help=f"Claude model for web research (default: {RESEARCH_MODEL})",
    )
    p.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Verbose logging",
    )
    return p.parse_args()


def resolve_inquiry(inquiry_arg: str) -> str:
    """If the inquiry argument is a file path, read it; otherwise return as-is."""
    path = Path(inquiry_arg)
    if path.is_file() and path.suffix in (".txt", ".md"):
        LOG.info("Reading inquiry from file: %s", path)
        return path.read_text(encoding="utf-8").strip()
    return inquiry_arg


def main():
    args = parse_args()
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%H:%M:%S",
    )

    # --- Validate inputs ---
    bio_path = Path(args.bio)
    if not bio_path.is_file():
        LOG.error("Bio file not found: %s", bio_path)
        sys.exit(1)

    profile_path = Path(args.company_profile)
    if not profile_path.is_file():
        LOG.error("Company profile not found: %s", profile_path)
        sys.exit(1)

    inquiry = resolve_inquiry(args.inquiry)
    if not inquiry:
        LOG.error("Inquiry text is empty.")
        sys.exit(1)

    company = args.company
    slug = slugify(company)

    # Output directory
    out_dir = Path(args.output_dir) if args.output_dir else SCRIPT_DIR / "output" / slug
    out_dir.mkdir(parents=True, exist_ok=True)

    # --- Init API client ---
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        LOG.error("ANTHROPIC_API_KEY not set. Export it or add to .env file.")
        sys.exit(1)
    client = anthropic.Anthropic(api_key=api_key)

    # --- Extract bio content ---
    LOG.info("Extracting bio from: %s", bio_path)
    original_bio = extract_bio_text(bio_path)
    LOG.info("Bio text: %d chars", len(original_bio))

    with tempfile.TemporaryDirectory() as tmp:
        images = extract_bio_images(bio_path, Path(tmp))
        LOG.info("Extracted %d images from bio", len(images))

        # --- Load Synaptiq profile ---
        synaptiq_profile = profile_path.read_text(encoding="utf-8")

        # --- Step 1: Research the company ---
        company_profile_md = research_company(
            client, company, inquiry, args.research_model
        )

        # --- Step 2: Generate tailored content ---
        bio_data = generate_tailored_bio(
            client, original_bio, company_profile_md, synaptiq_profile,
            inquiry, company, args.model
        )

        # --- Step 3: Generate prep doc (uses other outputs as input) ---
        prep_doc_md = generate_prep_doc(
            client, company_profile_md, synaptiq_profile, inquiry, company,
            bio_data["discovery_questions"], bio_data["tailoring_notes"],
            args.model,
        )

        # --- Step 4: Write all outputs ---
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 1. Tailored bio .docx
        docx_path = out_dir / f"Sklarew Bio - {company}.docx"
        build_docx(bio_data, images, docx_path)

        # 2. Tailoring specifics .md
        tailoring_path = out_dir / f"tailoring-specifics-{slug}_{ts}.md"
        write_tailoring_notes(bio_data["tailoring_notes"], company, tailoring_path)

        # 3. Company profile .md
        profile_out_path = out_dir / f"company-profile-{slug}_{ts}.md"
        write_company_profile(company_profile_md, company, profile_out_path)

        # 4. Discovery questions .md
        questions_path = out_dir / f"discovery-questions-{slug}_{ts}.md"
        write_discovery_questions(
            bio_data["discovery_questions"], company, questions_path
        )

        # 5. Consultant prep brief .md
        prep_path = out_dir / f"prep-brief-{slug}_{ts}.md"
        write_prep_doc(prep_doc_md, company, prep_path)

    # --- Summary ---
    LOG.info("--- Done ---")
    LOG.info("Company: %s", company)
    LOG.info("Outputs in: %s", out_dir)
    LOG.info("  1. Bio:       %s", docx_path.name)
    LOG.info("  2. Tailoring: %s", tailoring_path.name)
    LOG.info("  3. Profile:   %s", profile_out_path.name)
    LOG.info("  4. Questions: %s", questions_path.name)
    LOG.info("  5. Prep:      %s", prep_path.name)


if __name__ == "__main__":
    main()
