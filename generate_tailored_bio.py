#!/usr/bin/env python3
"""Generate a tailored bio document for Chick-fil-A inquiry."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# --- Logo ---
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = logo_para.add_run()
logo_run.add_picture('/tmp/bio_images/word/media/image1.png', width=Inches(3.5))
logo_para.space_after = Pt(4)

# --- Divider line ---
divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
div_run = divider.add_run('_' * 70)
div_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
div_run.font.size = Pt(8)
divider.space_after = Pt(12)

# --- Headshot + Name/Title block using table for side-by-side layout ---
table = doc.add_table(rows=1, cols=2)
table.autofit = True
table.columns[0].width = Inches(1.8)
table.columns[1].width = Inches(5.0)

# Remove table borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        # Set no borders
        borders = '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            borders += f'<w:{border} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        borders += '</w:tcBorders>'
        from lxml import etree
        tcPr.append(etree.fromstring(borders))

# Headshot cell
cell_img = table.cell(0, 0)
cell_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
img_run = cell_img.paragraphs[0].add_run()
img_run.add_picture('/tmp/bio_images/word/media/image2.jpg', width=Inches(1.5))

# Name/title cell
cell_text = table.cell(0, 1)
cell_text.vertical_alignment = 1  # CENTER

# Clear default paragraph
cell_text.paragraphs[0].text = ''
name_run = cell_text.paragraphs[0].add_run('Stephen Sklarew')
name_run.bold = True
name_run.font.size = Pt(20)
name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
cell_text.paragraphs[0].space_after = Pt(2)

title_para = cell_text.add_paragraph()
title_run = title_para.add_run('CEO & Co-Founder, Synaptiq')
title_run.font.size = Pt(13)
title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
title_run.italic = True
title_para.space_after = Pt(4)

tagline_para = cell_text.add_paragraph()
tagline_run = tagline_para.add_run('AI Strategy & Workforce Transformation  |  "The Humankind of AI"')
tagline_run.font.size = Pt(10)
tagline_run.font.color.rgb = RGBColor(0x66, 0x88, 0xAA)
tagline_para.space_after = Pt(0)

# --- Spacing ---
spacer = doc.add_paragraph()
spacer.space_before = Pt(6)
spacer.space_after = Pt(2)

# --- Succinct Bio ---
heading1 = doc.add_paragraph()
h1_run = heading1.add_run('Overview')
h1_run.bold = True
h1_run.font.size = Pt(13)
h1_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
heading1.space_after = Pt(6)

succinct = doc.add_paragraph()
succinct_text = (
    "Stephen Sklarew is a seasoned executive with over 30 years of experience at the intersection of "
    "business strategy, technology, and organizational transformation. As CEO and co-founder of Synaptiq, "
    "an AI strategy and solutions firm founded in 2015, Stephen specializes in helping values-driven "
    "organizations navigate the evolving AI landscape with intention and integrity. His expertise in "
    "AI-augmented workforce design, capability mapping, and future-state planning enables leaders to "
    "capture efficiency gains from automation while preserving organizational culture, talent pipelines, "
    "and strategic capability. Through Synaptiq's proprietary AIQ\u00ae methodology, Stephen guides "
    "organizations through responsible AI readiness assessment, workforce scenario modeling, and "
    "implementation planning\u2014ensuring that technology adoption strengthens rather than disrupts "
    "the human foundations of high-performing teams."
)
s_run = succinct.add_run(succinct_text)
s_run.font.size = Pt(11)
succinct.space_after = Pt(12)

# --- Extended Bio ---
heading2 = doc.add_paragraph()
h2_run = heading2.add_run('Background & Expertise')
h2_run.bold = True
h2_run.font.size = Pt(13)
h2_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
heading2.space_after = Pt(6)

para1 = doc.add_paragraph()
p1_text = (
    "Stephen Sklarew is a seasoned business and technology executive with over 30 years of experience "
    "driving innovation across startups and publicly traded companies. His expertise spans strategy and "
    "innovation, product development and management, and operations, with deep focus on Governance "
    "(Project Lifecycle Management and Analytics Lifecycle Management), Technology Management, and "
    "Product Orientation. At Synaptiq, Stephen focuses on helping organizations\u2014particularly those "
    "with strong cultural identities and deliberate approaches to change\u2014leverage AI and data to "
    "unlock business value while safeguarding the people-centered principles that define them."
)
para1.add_run(p1_text).font.size = Pt(11)
para1.space_after = Pt(8)

para2 = doc.add_paragraph()
p2_text = (
    "Stephen brings particular depth in guiding organizations through the workforce implications of AI "
    "adoption. He has advised leadership teams on how to redesign performance and insights functions as "
    "automation matures\u2014helping them develop scenario models for role evolution, capability mapping "
    "frameworks for existing teams, and strategic restructuring plans that prioritize talent development "
    "over workforce reduction. His approach recognizes that the most impactful AI transformations happen "
    "when organizations slow down to plan thoughtfully, aligning technology adoption with their values, "
    "growth trajectory, and long-term workforce strategy."
)
para2.add_run(p2_text).font.size = Pt(11)
para2.space_after = Pt(8)

para3 = doc.add_paragraph()
p3_text = (
    "Prior to founding Synaptiq, Stephen held senior positions at Sage Software, CEB (now Gartner), "
    "and Ernst & Young LLP, as well as several venture-backed startups including CircleBack "
    "(now SalesIntel), a SaaS company he co-founded in 2008 that leveraged big data and machine learning. "
    "This blend of enterprise consulting, startup agility, and hands-on technology leadership gives him a "
    "rare ability to translate between executive strategy and technical execution\u2014meeting organizations "
    "where they are and guiding them toward where they need to be."
)
para3.add_run(p3_text).font.size = Pt(11)
para3.space_after = Pt(8)

para4 = doc.add_paragraph()
p4_text = (
    "Stephen has spearheaded technology and organizational transformation initiatives across diverse "
    "sectors including food and beverage, healthcare, legal, construction, financial services, and the "
    "public sector. He is deeply committed to building solutions that catalyze human ingenuity rather "
    "than replace it\u2014a philosophy captured in Synaptiq's tagline, \"The Humankind of AI.\" Under his "
    "leadership, Synaptiq developed its proprietary AIQ\u00ae methodology, a comprehensive framework that "
    "guides clients from AI readiness assessment through strategy development, workforce planning, and "
    "implementation. This approach integrates Analytics Strategy, Data Strategy, and organizational change "
    "management to ensure that AI adoption creates lasting business value while maintaining the stability "
    "and developmental culture that great organizations are built on."
)
para4.add_run(p4_text).font.size = Pt(11)
para4.space_after = Pt(8)

para5 = doc.add_paragraph()
p5_text = (
    "Stephen is a thought leader in AI and its evolving role in business and workforce design, "
    "frequently sharing insights as an author, speaker, and guest lecturer. He holds a BS in Biology "
    "from Virginia Tech."
)
para5.add_run(p5_text).font.size = Pt(11)
para5.space_after = Pt(12)

# --- Relevant Experience Highlights ---
heading3 = doc.add_paragraph()
h3_run = heading3.add_run('Relevant Experience')
h3_run.bold = True
h3_run.font.size = Pt(13)
h3_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
heading3.space_after = Pt(6)

highlights = [
    ("AI & Workforce Strategy", "Guided organizations in designing future-state workforce models that capture automation efficiencies while preserving culture and talent development pathways"),
    ("Performance & Analytics Team Design", "Advised on restructuring insights and analytics functions to evolve alongside AI capabilities, including capability mapping and role redesign frameworks"),
    ("Food & Beverage Industry", "Delivered AI-powered solutions for food service operations including machine vision, supply chain optimization, and operational intelligence"),
    ("AIQ\u00ae Readiness Methodology", "Developed and deployed a proprietary AI maturity assessment across four dimensions\u2014Organizational Foundations, Product Lifecycle, Data Infrastructure, and AI/ML Capabilities\u2014helping organizations benchmark readiness and plan adoption responsibly"),
    ("Enterprise Data Strategy", "Built unified data management architectures (DataLake) enabling organizations to consolidate siloed data for analytics and decision-making at scale"),
    ("Organizational Change Management", "Supported leadership teams in managing the human side of technology transformation, ensuring stakeholder alignment and cultural continuity"),
]

for title, desc in highlights:
    bullet = doc.add_paragraph()
    bullet.style = 'List Bullet'
    title_run = bullet.add_run(f'{title}: ')
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    desc_run = bullet.add_run(desc)
    desc_run.font.size = Pt(10.5)
    bullet.space_after = Pt(4)

# Save
output_path = '/Users/stephensklarew/Development/Scripts/synaptiq-iia-inquiry-analysis/Sklarew Bio - Chick-fil-A.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
